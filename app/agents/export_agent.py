# ============================================================
# app/agents/export_agent.py
# ============================================================
# Export Agent — Phase 3 (Post-HITL Approval)
#
# Converts the final Markdown report into:
#   - PDF  : Professional branded document (ReportLab)
#   - Word : Fully editable document (python-docx)
#
# Both formats include:
#   - Title page with doctor approval stamp
#   - Section headings and bullet points
#   - Numbered source citations
#   - Medical disclaimer on every page
#
# Flow:
#   1. Read final_report from state
#   2. Parse Markdown into sections
#   3. Generate PDF and/or Word file
#   4. Save file paths back to state
# ============================================================

import os
import uuid
from datetime import datetime
from typing import Optional

# ── PDF Generation ────────────────────────────────────────────
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer,
    HRFlowable, PageBreak
)

# ── Word Generation ───────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.config import config
from app.graph.state import ResearchState


# ── Setup Export Directory ────────────────────────────────────

def ensure_export_dir() -> str:
    """
    Create the exports directory if it does not exist.

    Returns:
        Path to the exports directory.
    """
    export_dir = config.EXPORT_CACHE_DIR
    os.makedirs(export_dir, exist_ok=True)
    return export_dir

# ── Parse Markdown ────────────────────────────────────────────

def parse_markdown_sections(report: str) -> list:
    """
    Parse a Markdown report string into a list of sections.

    Each section is a dict with:
        type:  'h1', 'h2', 'bullet', 'text', 'divider'
        text:  The content string

    Args:
        report: Markdown report string.

    Returns:
        List of section dicts.
    """
    sections = []
    lines = report.split("\n")
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line — add spacing
            sections.append({"type": "space", "text": ""})
            
        elif line.startswith("# "):
            # H1 heading
            sections.append({"type": "h1", "text": line[2:].strip()})
            
        elif line.startswith("## "):
             # H2 heading
            sections.append({"type": "h2", "text": line[3:].strip()})

        elif line.startswith("### "):
            # H3 heading
            sections.append({"type": "h3", "text": line[4:].strip()})
            
        elif line.startswith("- ") or line.startswith("* "):
            # Bullet point
            sections.append({"type": "bullet", "text": line[2:].strip()})

        elif line.startswith("---"):
            # Horizontal divider
            sections.append({"type": "divider", "text": ""})

        elif line.startswith("**") and line.endswith("**"):
            # Bold text
            sections.append({"type": "bold", "text": line.strip("*")})

        else:
            # Regular paragraph text
            sections.append({"type": "text", "text": line})
    
    return sections

# ── PDF Generation ────────────────────────────────────────────

def generate_pdf(report: str, query: str, confidence: int,
    hitl_comments: str, session_id: str) -> str:
    """
    Generate a professional PDF report using ReportLab.

    Args:
        report:        Final Markdown report string.
        query:         Original research query.
        confidence:    Confidence score 0-100.
        hitl_comments: Doctor's approval comments.
        session_id:    Unique session identifier.

    Returns:
        File path to the generated PDF.
    """
    export_dir = ensure_export_dir()
    filename = f"report_{session_id[:8]}_{datetime.now().strftime('%Y%m%d')}.pdf"
    filepath = os.path.join(export_dir, filename)
    
    # ── Document Setup ────────────────────────────────────────
    doc = SimpleDocTemplate(
        filepath,
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    
     # ── Styles ────────────────────────────────────────────────
    styles = getSampleStyleSheet()

    # Custom styles for our report
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=24,
        textColor=colors.HexColor("#1F3864"),
        spaceAfter=12,
        fontName="Helvetica-Bold"
    )
    
    h1_style = ParagraphStyle(
        "CustomH1",
        parent=styles["Heading1"],
        fontSize=16,
        textColor=colors.HexColor("#1F3864"),
        spaceBefore=16,
        spaceAfter=8,
        fontName="Helvetica-Bold"
    )
    
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#2E75B6"),
        spaceBefore=12,
        spaceAfter=6,
        fontName="Helvetica-Bold"
    )
    
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#404040"),
        spaceAfter=8,
        leading=16,
        fontName="Helvetica"
    )
    
    bullet_style = ParagraphStyle(
        "CustomBullet",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#404040"),
        spaceAfter=6,
        leftIndent=20,
        fontName="Helvetica"
    )
    
    disclaimer_style = ParagraphStyle(
        "Disclaimer",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.HexColor("#888888"),
        spaceAfter=6,
        fontName="Helvetica-Oblique"
    )
    
    # ── Build Content ─────────────────────────────────────────
    content = []

    # Title page
    content.append(Spacer(1, 0.5*inch))
    content.append(Paragraph("MediResearch AI", title_style))
    content.append(Paragraph("Medical Research Report", h2_style))
    content.append(HRFlowable(width="100%", thickness=2,
                               color=colors.HexColor("#2E75B6")))
    content.append(Spacer(1, 0.2 * inch))
    
    # Report metadata
    content.append(Paragraph(f"<b>Query:</b> {query}", body_style))
    content.append(Paragraph(
        f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')}", body_style))
    content.append(Paragraph(
        f"<b>Confidence Score:</b> {confidence}/100", body_style))
    content.append(Paragraph(
        f"<b>Doctor Approval:</b> ✓ Approved", body_style))
    
    # Doctor comments if any
    if hitl_comments and hitl_comments != "No additional comments.":
        content.append(Paragraph(
            f"<b>Doctor Notes:</b> {hitl_comments}", body_style))

    content.append(Spacer(1, 0.3 * inch))
    content.append(HRFlowable(width="100%", thickness=1,
                               color=colors.HexColor("#CCCCCC")))
    content.append(PageBreak())

    # Report body — parse and render each section
    sections = parse_markdown_sections(report)
    
    for section in sections:
        stype = section["type"]
        text = section["text"]
        
        if not text and stype not in ["space", "divider"]:
            continue
        
        if stype == "h1":
            content.append(Paragraph(text, h1_style))
            
        elif stype == "h2":
            content.append(Paragraph(text, h2_style))
            
        elif stype == "h3":
            content.append(Paragraph(f"<b>{text}</b>", body_style))
            
        elif stype == "bullet":
            content.append(Paragraph(f"• {text}", bullet_style))
            
        elif stype  == "bold":
             content.append(Paragraph(f"<b>{text}</b>", body_style))

        elif stype == "text":
            content.append(Paragraph(text, body_style))

        elif stype == "divider":
            content.append(HRFlowable(
                width="100%", thickness=1,
                color=colors.HexColor("#CCCCCC")
            ))

        elif stype == "space":
            content.append(Spacer(1, 0.1 * inch))

    # Disclaimer footer
    content.append(Spacer(1, 0.3 * inch))
    content.append(HRFlowable(width="100%", thickness=1,
                               color=colors.HexColor("#CCCCCC")))
    
    content.append(Paragraph(
        "⚠️ MEDICAL DISCLAIMER: This report is generated by an AI system "
        "for research and educational purposes only. It does not constitute "
        "medical advice. Always consult a qualified medical professional "
        "for clinical decisions.",
        disclaimer_style
    ))

    # Build the PDF
    doc.build(content)
    return filepath


# ── Word Generation ───────────────────────────────────────────

def generate_word(report: str, query: str, confidence: int,
    hitl_comments: str, session_id: str) -> str:
    """
    Generate a professional Word document using python-docx.

    Args:
        report:        Final Markdown report string.
        query:         Original research query.
        confidence:    Confidence score 0-100.
        hitl_comments: Doctor's approval comments.
        session_id:    Unique session identifier.

    Returns:
        File path to the generated Word document.
    """
    export_dir = ensure_export_dir()
    filename = f"report_{session_id[:8]}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(export_dir, filename)
    
    # ── Create Document ───────────────────────────────────────
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ── Title Page ────────────────────────────────────────────
    title = doc.add_heading("MediResearch AI", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    subtitle = doc.add_heading("Medical Research Report", level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    
    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    # Fill metadata rows
    meta = [
        ("Query", query),
        ("Date", datetime.now().strftime("%B %d, %Y")),
        ("Confidence Score", f"{confidence}/100"),
        ("Doctor Approval", "✓ Approved"),
    ]
    for i, (label, value) in enumerate(meta):
        row = table.rows[i]
        # Label cell — bold navy
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        run        = label_para.add_run(label)
        run.bold   = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        # Value cell
        value_cell = row.cells[1]
        value_cell.paragraphs[0].add_run(value)

    # Doctor comments
    if hitl_comments and hitl_comments != "No additional comments.":
        doc.add_paragraph()
        p    = doc.add_paragraph()
        run  = p.add_run("Doctor Notes: ")
        run.bold = True
        p.add_run(hitl_comments)

    doc.add_page_break()
    
    # ── Report Body ───────────────────────────────────────────
    sections = parse_markdown_sections(report)

    for section in sections:
        stype = section["type"]
        text  = section["text"]

        if not text and stype not in ["space", "divider"]:
            continue

        if stype == "h1":
            h = doc.add_heading(text, level=1)
            h.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

        elif stype == "h2":
            h = doc.add_heading(text, level=2)
            h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            
        elif stype == "h3":
            p    = doc.add_paragraph()
            run  = p.add_run(text)
            run.bold = True

        elif stype == "bullet":
            doc.add_paragraph(text, style="List Bullet")

        elif stype == "bold":
            p    = doc.add_paragraph()
            run  = p.add_run(text)
            run.bold = True

        elif stype == "text":
            doc.add_paragraph(text)

        elif stype == "space":
            doc.add_paragraph()
            
    # ── Disclaimer ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_heading("⚠️ Medical Disclaimer", level=2)
    disclaimer = doc.add_paragraph(
        "This report is generated by an AI system for research and "
        "educational purposes only. It does not constitute medical advice. "
        "Always consult a qualified medical professional for clinical decisions."
    )
    disclaimer.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    disclaimer.runs[0].italic = True

    # Save the document
    doc.save(filepath)
    return filepath

# ── Main Export Agent ─────────────────────────────────────────

def run_export_agent(state: ResearchState) -> ResearchState:
    """
    Main Export Agent — called by LangGraph after HITL approval.

    Generates both PDF and Word versions of the final report.

    Reads from state:  final_report, query, confidence_score,
                       hitl_comments, session_id
    Writes to state:   export_pdf_path, export_word_path

    Args:
        state: Current ResearchState from LangGraph.

    Returns:
        Updated state with export file paths.
    """
    print("📤 Export Agent running...")
    
    final_report = state.get("final_report", "")
    query = state.get("query", "Unknown Query")
    confidence = state.get("confidence_score", 0)
    hitl_comments = state.get("hitl_comments", "No additional comments.")
    session_id = state.get("session_id", str(uuid.uuid4()))
    
    if not final_report:
        print("   ⚠️  No final report found — skipping export")
        return state
    
    pdf_path = None
    word_path = None
    
    # ── Generate PDF ──────────────────────────────────────────
    try:
        print("   Generating PDF...")
        pdf_path = generate_pdf(
            report=final_report,
            query=query,
            confidence=confidence,
            hitl_comments=hitl_comments,
            session_id=session_id
        )
        print(f"   ✅ PDF saved: {pdf_path}")
        
    except Exception as e:
        print(f"   ❌ PDF generation failed: {e}")
        
    # ── Generate Word ─────────────────────────────────────────
    try:
        print("   Generating Word document...")
        word_path = generate_word(
            report=final_report,
            query=query,
            confidence=confidence,
            hitl_comments=hitl_comments,
            session_id=session_id
        )
        print(f"   ✅ Word saved: {word_path}")

    except Exception as e:
        print(f"   ❌ Word generation failed: {e}")

    return {
        **state,
        "export_pdf_path":  pdf_path,
        "export_word_path": word_path
    }
